"""Normalize CV bullet/link indentation for a clean, consistent look."""
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Twips, RGBColor

SRC = Path(r"d:\Projects\CV\CV_Nuno_Tamada_EN_links.docx")
OUT = Path(r"d:\Projects\CV\CV_Nuno_Tamada_EN_clean.docx")

# Match existing Konek list indent (from inspection)
LIST_LEFT = Twips(161290 // 20)  # python-docx uses EMU via Twips? Actually Pt/Inches/Twips/Cm/Mm/Emu
# 161290 EMU — docx.shared.Emu
from docx.shared import Emu

LIST_LEFT_EMU = Emu(161290)
LIST_HANGING_EMU = Emu(125095)  # first_line_indent negative
URL_LEFT_EMU = Emu(161290)  # align with bullet TEXT, not the bullet glyph


def set_run_font(run, size_pt=9.5, bold=None, name="Calibri", color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def clear_runs_keep_one(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]:
            r.text = ""
        return paragraph.runs[0]
    return paragraph.add_run(text)


def make_body_paragraph(paragraph, text, size=9.5):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = LIST_LEFT_EMU
    paragraph.paragraph_format.first_line_indent = Emu(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    run = clear_runs_keep_one(paragraph, text)
    set_run_font(run, size_pt=size, bold=False, name="Calibri")
    return run


def make_bullet_paragraph(paragraph, text, size=9.5):
    """Hanging indent: bullet hangs left, wrapped lines align with first text column."""
    paragraph.style = paragraph.part.document.styles["List Paragraph"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = LIST_LEFT_EMU
    paragraph.paragraph_format.first_line_indent = Emu(-int(LIST_HANGING_EMU))
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15

    # Ensure Word list bullet exists; if text already starts with • strip it
    clean = text.lstrip("•").strip()
    # Prefer native list bullet — remove leading bullet chars from text
    run = clear_runs_keep_one(paragraph, clean)
    set_run_font(run, size_pt=size, bold=False, name="Calibri")
    return run


def make_link_paragraph(paragraph, label, size=9.5):
    """URL aligned with bullet text column (same left indent, no hanging)."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.style = paragraph.part.document.styles["Normal"]
    paragraph.paragraph_format.left_indent = URL_LEFT_EMU
    paragraph.paragraph_format.first_line_indent = Emu(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    run = clear_runs_keep_one(paragraph, label)
    set_run_font(run, size_pt=size, bold=False, name="Calibri", color=RGBColor(0x0F, 0x5E, 0x8C))
    run.underline = True
    return run


def main():
    d = Document(str(SRC))

    # Konek Market bullets 15-20 + link 21
    konek_bullets = range(15, 21)
    for i in konek_bullets:
        t = d.paragraphs[i].text.lstrip("•").strip()
        make_bullet_paragraph(d.paragraphs[i], t)

    make_link_paragraph(
        d.paragraphs[21],
        "play.google.com/store/apps/details?id=com.konekmarket",
    )

    # Drive Test 23-27: currently fake bullets in text — convert to real list style
    for i in range(23, 28):
        t = d.paragraphs[i].text.lstrip("•").strip()
        make_bullet_paragraph(d.paragraphs[i], t)

    # Teaching assistant 29-30 — make bullets for consistency if not already
    for i in (29, 30):
        t = d.paragraphs[i].text.lstrip("•").strip()
        make_bullet_paragraph(d.paragraphs[i], t)

    # Project bullets Edu/LearnQuest
    for i in (34, 35, 39, 40):
        t = d.paragraphs[i].text.lstrip("•").strip()
        make_bullet_paragraph(d.paragraphs[i], t)

    # Project links aligned like Konek link
    for i, label in [
        (36, "edu-app-ec176.web.app"),
        (41, "learnquest-lms-59vf.vercel.app"),
        (45, "lumen-five-umber.vercel.app"),
        (49, "play.google.com/store/apps/details?id=com.konekmarket"),
    ]:
        make_link_paragraph(d.paragraphs[i], label)

    # Lumen description body indent same as other project desc
    if d.paragraphs[44].text.strip():
        make_body_paragraph(d.paragraphs[44], d.paragraphs[44].text.strip())
        d.paragraphs[44].paragraph_format.left_indent = Emu(0)  # projects often full width
        d.paragraphs[48].paragraph_format.left_indent = Emu(0)

    d.save(str(OUT))
    print("saved", OUT)

    # verify indents
    d2 = Document(str(OUT))
    for i in (15, 16, 20, 21, 23, 36, 45):
        pf = d2.paragraphs[i].paragraph_format
        print(
            i,
            "left",
            int(pf.left_indent or 0),
            "first",
            int(pf.first_line_indent or 0),
            "=>",
            d2.paragraphs[i].text[:60],
        )


if __name__ == "__main__":
    main()
