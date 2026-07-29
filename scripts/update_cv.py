from docx import Document
from pathlib import Path

src = Path(r"c:\Users\Nuno Tamada\Downloads\Nuno_Tamada_CV_Revised.docx")
dst = Path(r"c:\Users\Nuno Tamada\Downloads\Nuno_Tamada_CV_Revised_Lumen.docx")
d = Document(str(src))

TITLE = "Lumen — AI Knowledge Desk\tNext.js · TypeScript · RAG · Groq"
DESC = (
    "RAG knowledge desk: ingest MD/PDF, hybrid retrieval (vector + BM25), "
    "streaming answers with source citations, map-reduce summarization, and "
    "golden-set evaluation (10 docs · 11 chunks · chat ~0.8s · Recall@4 = 100%). "
    "Live demo (read-only seed) on Vercel."
)
SUMMARY = (
    "Web, mobile, and AI developer experienced in React Native, React/Next.js, "
    "and Flutter-based learning apps. Built and shipped mobile features for Konek Market "
    "(published on Google Play), developed full-stack web products with Next.js and Prisma, "
    "and shipped Lumen — a RAG knowledge desk with hybrid retrieval, streaming answers, "
    "citations, and Recall@K evaluation. Comfortable with REST API integration, debugging, "
    "documentation-minded delivery, and AI-assisted workflows (Claude, Copilot) for rapid prototyping."
)


def set_para_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


changed = []
for p in d.paragraphs:
    t = p.text.strip()

    if t == "Web & Mobile Developer":
        set_para_text(p, "Web, Mobile & AI Developer")
        changed.append("title")
        continue

    if t.startswith("Web & mobile developer experienced"):
        set_para_text(p, SUMMARY)
        changed.append("summary")
        continue

    if t.startswith("Backend & Data:"):
        set_para_text(
            p,
            "Backend & Data:  Firebase, Prisma, PostgreSQL, REST APIs, Axios, Groq (LLM APIs)",
        )
        changed.append("backend")
        continue

    if t.startswith("Practices:"):
        set_para_text(
            p,
            "Practices:  Git, Debugging, UI implementation from design, API integration, RAG pipelines, retrieval evaluation",
        )
        changed.append("practices")
        continue

    if t.startswith("Lumen") and "Knowledge Desk" in t:
        set_para_text(p, TITLE)
        changed.append("lumen-title")
        continue

    if (
        "hybrid retrieval" in t.lower()
        and "streaming" in t.lower()
        and "learnquest" not in t.lower()
        and not t.lower().startswith("web")
    ):
        set_para_text(p, DESC)
        changed.append("lumen-desc")
        continue

d.save(str(dst))
print("saved:", dst)
print("changed:", changed)

d2 = Document(str(dst))
for p in d2.paragraphs:
    if any(
        k in p.text
        for k in ("Lumen", "Recall@4", "AI Developer", "Groq", "RAG pipelines")
    ):
        print(">>", p.text)
